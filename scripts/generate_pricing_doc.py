# -*- coding: utf-8 -*-
# File: scripts/generate_pricing_doc.py
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_pricing_document():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(51, 51, 51) # Charcoal
    style_normal.paragraph_format.line_spacing = 1.3
    style_normal.paragraph_format.space_after = Pt(6)

    # Custom Colors
    PRIMARY_COLOR = RGBColor(31, 78, 121)    # Deep Steel Blue
    SECONDARY_COLOR = RGBColor(79, 93, 115)  # Muted Indigo Slate
    HIGHLIGHT_COLOR = RGBColor(184, 15, 10)  # Dark Red/Coral for risks
    MUTED_GRAY = RGBColor(120, 120, 120)     # Gray for subtitles

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    title_run = title_p.add_run("BÁO CÁO PHÂN TÍCH TÀI CHÍNH & CHIẾN LƯỢC DOANH THU\nDỰ ÁN HỆ SINH THÁI SMARTLIFE")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(40)
    sub_run = subtitle_p.add_run("Phân Tích Chi Phí Token AI, Phương Án Phòng Ngừa Rủi Ro Bù Lỗ,\nMô Hình Định Giá Doanh Thu Mới & Định Vị Thị Trường TAM - SAM - SOM")
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = SECONDARY_COLOR

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_before = Pt(80)
    info_p.paragraph_format.space_after = Pt(40)
    
    author_run = info_p.add_run("Đơn vị thực hiện: Ban Quản trị Dự án SmartLife\nChuyên gia tư vấn: AI Technology & Business Advisor\nThời gian báo cáo: Tháng 06, 2026\nPhiên bản tài liệu: v2.1 (Cập nhật sau tái cấu trúc)")
    author_run.font.size = Pt(11)
    author_run.font.color.rgb = SECONDARY_COLOR

    doc.add_page_break()

    # --- Section: Agenda ---
    agenda_title = doc.add_paragraph()
    agenda_run = agenda_title.add_run("MỤC LỤC CHƯƠNG TRÌNH (AGENDA)")
    agenda_run.font.size = Pt(16)
    agenda_run.font.bold = True
    agenda_run.font.color.rgb = PRIMARY_COLOR
    agenda_title.paragraph_format.space_after = Pt(12)

    agenda_items = [
        ("PHẦN I: TÁI CẤU TRÚC GIÁ & BẢN ĐỒ PHÂN KHÚC DỊCH VỤ MỚI", 3),
        ("  1. Phân cấp tính năng & Chặn truy cập AI của Gói Free", 3),
        ("  2. Đồng bộ biểu giá mới (1 tháng, 3 tháng, 12 tháng, Trọn đời)", 3),
        ("  3. Cơ chế kích hoạt và tiêu dùng của Gói AI Boost Packs", 4),
        ("PHẦN II: PHÂN TÍCH CHI PHÍ TOKEN AI & PHƯƠNG ÁN PHÒNG NGỪA BÙ LỖ", 5),
        ("  1. Chi phí API Google Gemini 2.5 Flash thực tế", 5),
        ("  2. Bài toán chi phí trên mỗi người dùng Pro hoạt động tối đa", 5),
        ("  3. Đánh giá tỷ suất lợi nhuận biên của Gói Boost Packs", 6),
        ("PHẦN III: QUẢN TRỊ RỦI RO CHI PHÍ & TẤM LÁ CHẮN BẢO VỆ HỆ THỐNG", 7),
        ("  1. Chốt chặn Server-Side Quota (Giới hạn Requests & Tokens)", 7),
        ("  2. Bảo mật API Keys & Xác thực Token JWT Supabase", 7),
        ("  3. Biện pháp tối ưu hóa Context giảm thiểu lạm dụng", 8),
        ("PHẦN IV: ĐỊNH VỊ THỊ TRƯỜNG THEO MÔ HÌNH TAM - SAM - SOM", 9),
        ("  1. TAM (Total Addressable Market) - Thị trường tổng thể", 9),
        ("  2. SAM (Serviceable Addressable Market) - Thị trường tiếp cận", 9),
        ("  3. SOM (Serviceable Obtainable Market) - Thị trường mục tiêu thực tế", 10),
        ("  4. Dự phóng Doanh thu & Điểm hòa vốn", 11)
    ]

    for item, page in agenda_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_name = p.add_run(item.ljust(85, '.'))
        run_name.font.size = Pt(11)
        if "PHẦN" in item:
            run_name.font.bold = True
        run_page = p.add_run(f" Trang {page}")
        run_page.font.size = Pt(11)
        if "PHẦN" in item:
            run_page.font.bold = True

    doc.add_page_break()

    # --- Section 1 ---
    h1 = doc.add_paragraph()
    r = h1.add_run("PHẦN I: TÁI CẤU TRÚC GIÁ & BẢN ĐỒ PHÂN KHÚC DỊCH VỤ MỚI")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)

    p1 = doc.add_paragraph(
        "Nhằm tối ưu hóa nguồn thu doanh nghiệp và cân bằng chi phí hạ tầng máy chủ AI, dự án SmartLife đã thực hiện một đợt nâng cấp biểu giá toàn diện và tái cấu trúc sâu phân cấp dịch vụ. Thay đổi lớn nhất trong bản cập nhật này là việc chuyển giao AI từ một tính năng mở rộng thành đặc quyền thương mại cao cấp hoàn toàn nằm sau cổng thanh toán (Pro Gate)."
    )

    h2 = doc.add_paragraph()
    r = h2.add_run("1. Phân cấp tính năng & Chặn truy cập AI của Gói Free")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Người dùng thuộc gói Free (Miễn phí) sẽ không còn quyền sử dụng bất kỳ tính năng AI nào của hệ thống. Bản đồ phân cấp tính năng hiện tại được phân chia cụ thể như sau:"
    )

    # Table 1: Feature Matrix
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tính năng / Phân hệ'
    hdr_cells[1].text = 'Gói Free (Miễn phí)'
    hdr_cells[2].text = 'Gói Pro / Lifetime'
    
    # Format Table Header
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        cell.width = Inches(2.2)

    features = [
        ("Lịch trình, Todolist, Habit, Pomodoro", "Đầy đủ (Miễn phí hoàn toàn)", "Đầy đủ"),
        ("GPA Tracker & Lộ trình môn học", "Chỉ xem & nhập điểm thủ công", "Đầy đủ + Gợi ý lộ trình cải thiện"),
        ("Tài chính & Ví điện tử nâng cao", "Cơ bản (Giới hạn ví phụ)", "Đầy đủ (Không giới hạn ví & phân tích)"),
        ("AI Advisor (Cố vấn học tập/tài chính)", "KHÓA HOÀN TOÀN (Hiện Pro Gate)", "Đầy đủ (Hạn mức 600.000 tokens/tháng)"),
        ("Visual Board & Báo cáo tổng quan", "Khóa hoàn toàn", "Đầy đủ thông tin đa chiều")
    ]

    for idx, (f_name, free_status, pro_status) in enumerate(features):
        row_cells = table.rows[idx+1].cells
        row_cells[0].text = f_name
        row_cells[1].text = free_status
        row_cells[2].text = pro_status
        
        # Color formatting
        for c_idx, cell in enumerate(row_cells):
            cell.width = Inches(2.2)
            for p in cell.paragraphs:
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10.5)

    # Background styling hack for python-docx (applying colors to cells)
    # Background color helper
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    def set_cell_background(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Set Header color to Deep Steel Blue (1F4E79)
    for cell in table.rows[0].cells:
        set_cell_background(cell, "1F4E79")

    # Set stripe color for rows
    for r_idx in range(1, len(table.rows)):
        if r_idx % 2 == 0:
            for cell in table.rows[r_idx].cells:
                set_cell_background(cell, "F2F5F8")

    doc.add_paragraph() # Spacer

    h2 = doc.add_paragraph()
    r = h2.add_run("2. Đồng bộ biểu giá mới (1 tháng, 3 tháng, 12 tháng, Trọn đời)")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Nhằm tăng khả năng tiếp cận của người dùng đồng thời tối đa hóa Giá trị trọn đời của khách hàng (LTV), cấu trúc giá mới của gói Pro/Lifetime được cấu hình như sau:"
    )

    # Table 2: Pricing Matrix
    table2 = doc.add_table(rows=5, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Tên gói đăng ký'
    hdr2[1].text = 'Đơn giá thanh toán'
    hdr2[2].text = 'Quy đổi tháng'
    hdr2[3].text = 'Tỷ lệ tiết kiệm (%)'

    for cell in hdr2:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    pricing = [
        ("Gói Pro 1 Tháng", "59.000đ", "59.000đ / tháng", "Gói cơ sở (0%)"),
        ("Gói Pro 3 Tháng", "139.000đ", "46.333đ / tháng", "Tiết kiệm 21%"),
        ("Gói Pro 12 Tháng", "399.000đ", "33.250đ / tháng", "Tiết kiệm 44%"),
        ("Gói Lifetime (Trọn đời)", "899.000đ", "Sử dụng vĩnh viễn", "Giá trị tối ưu nhất")
    ]

    for idx, (p_name, price, monthly, save) in enumerate(pricing):
        row = table2.rows[idx+1].cells
        row[0].text = p_name
        row[1].text = price
        row[2].text = monthly
        row[3].text = save
        
        for c_idx, cell in enumerate(row):
            for p in cell.paragraphs:
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10.5)
        if (idx+1) % 2 == 0:
            for cell in row:
                set_cell_background(cell, "F2F5F8")

    doc.add_paragraph() # Spacer

    h2 = doc.add_paragraph()
    r = h2.add_run("3. Cơ chế kích hoạt và tiêu dùng của Gói AI Boost Packs")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Đối với các tài khoản Pro hoặc Lifetime sử dụng vượt quá dung lượng 600.000 tokens/tháng được định mức đi kèm, hệ thống sẽ tự động chuyển sang chế độ khóa hạn mức và hướng dẫn mua thêm các gói AI Boost Packs. Gói Boost Pack hoạt động theo các nguyên tắc kỹ thuật sau:"
    )

    bp_items = [
        ("Cơ chế tiêu thụ FIFO (First-In, First-Out):", " Token mua thêm từ các gói Boost chỉ được tiêu trừ sau khi người dùng đã tiêu dùng hết 100% dung lượng 600k tokens của tháng đó. Các gói mua trước sẽ được trừ trước."),
        ("Tính độc lập của gói đăng ký:", " Khi admin kích hoạt hóa đơn Boost Pack, hệ thống chỉ ghi nhận bổ sung số lượng token tương ứng vào bảng user_ai_boost mà hoàn toàn không ghi đè hoặc làm ảnh hưởng đến thời hạn gói Pro chính của tài khoản."),
        ("Thời hạn hiệu lực của Boost Pack:", " Các gói Boost S, M, L có thời hạn sử dụng độc lập lần lượt là 30, 60 và 90 ngày kể từ ngày kích hoạt thành công. Sau thời gian này, số token chưa sử dụng hết sẽ tự động hết hạn để tối ưu không gian cơ sở dữ liệu.")
    ]

    for title, desc in bp_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(title)
        run_t.font.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # --- Section 2 ---
    h1 = doc.add_paragraph()
    r = h1.add_run("PHẦN II: PHÂN TÍCH CHI PHÍ TOKEN AI & PHƯƠNG ÁN PHÒNG NGỪA BÙ LỖ")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)

    doc.add_paragraph(
        "Bảo đảm biên lợi nhuận ròng là ưu tiên hàng đầu của mô hình SaaS. Việc định lượng chi phí gọi API từ Google Gemini và đối soát với giá trị gói đăng ký của người dùng giúp SmartLife hoàn toàn tránh được rủi ro bù lỗ hạ tầng máy chủ."
    )

    h2 = doc.add_paragraph()
    r = h2.add_run("1. Chi phí API Google Gemini 2.5 Flash thực tế")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Gemini 2.5 Flash được lựa chọn là mô hình AI chủ lực cho phiên bản thương mại nhờ tốc độ xử lý vượt trội và đơn giá tối ưu. Đơn giá công bố từ nhà cung cấp Google AI Studio như sau:"
    )

    api_points = [
        ("Input Token (Đầu vào):", " 0.075 USD / 1.000.000 tokens (~1.912 VNĐ / triệu tokens)"),
        ("Output Token (Đầu ra):", " 0.300 USD / 1.000.000 tokens (~7.650 VNĐ / triệu tokens)"),
        ("Đơn giá trung bình ước tính:", " Dựa trên hành vi sử dụng thực tế của người dùng, tỷ lệ độ dài câu hỏi/ngữ cảnh đầu vào so với câu trả lời đầu ra là 80% Input và 20% Output. Từ đó, đơn giá trung bình cho mỗi triệu tokens tiêu thụ được tính toán là: (1.912 VNĐ * 0.8) + (7.650 VNĐ * 0.2) = 3.060 VNĐ / 1.000.000 tokens.")
    ]

    for title, desc in api_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(title)
        run_t.font.bold = True
        p.add_run(desc)

    h2 = doc.add_paragraph()
    r = h2.add_run("2. Bài toán chi phí trên mỗi người dùng Pro hoạt động tối đa")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Giả sử một người dùng Pro hoạt động cực kỳ tích cực và tiêu thụ hết 100% giới hạn token hàng tháng là 600.000 tokens. Chi phí biến đổi của người dùng này được tính như sau:"
    )

    doc.add_paragraph(
        "Chi phí API tối đa/tháng = 0.6 triệu tokens * 3.060 VNĐ/triệu tokens = 1.836 VNĐ / tháng.\n"
        "Đối soát với doanh thu thu được từ các gói đăng ký Pro tương ứng:"
    )

    # Table 3: Profit Matrix
    table3 = doc.add_table(rows=4, cols=4)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr3 = table3.rows[0].cells
    hdr3[0].text = 'Gói Pro đăng ký'
    hdr3[1].text = 'Doanh thu quy đổi/tháng'
    hdr3[2].text = 'Chi phí AI tối đa/tháng'
    hdr3[3].text = 'Biên lợi nhuận gộp (%)'

    for cell in hdr3:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    profits = [
        ("Gói Pro 1 Tháng", "59.000đ", "1.836đ", "96,9%"),
        ("Gói Pro 3 Tháng", "46.333đ", "1.836đ", "96,0%"),
        ("Gói Pro 12 Tháng", "33.250đ", "1.836đ", "94,5%")
    ]

    for idx, (p_name, rev, cost, margin) in enumerate(profits):
        row = table3.rows[idx+1].cells
        row[0].text = p_name
        row[1].text = rev
        row[2].text = cost
        row[3].text = margin
        
        for c_idx, cell in enumerate(row):
            for p in cell.paragraphs:
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10.5)
        if (idx+1) % 2 == 0:
            for cell in row:
                set_cell_background(cell, "F2F5F8")

    doc.add_paragraph() # Spacer

    doc.add_paragraph(
        "Nhận xét: Kịch bản tệ nhất khi người dùng Pro dùng hết 100% dung lượng đi kèm thì biên lợi nhuận gộp của sản phẩm vẫn đạt trên 94.5%. Đây là mức biên lợi nhuận cực kỳ lý tưởng đối với một mô hình SaaS, bảo đảm bù đắp dư thừa chi phí server vận hành, lưu trữ dữ liệu và marketing."
    )

    h2 = doc.add_paragraph()
    r = h2.add_run("3. Đánh giá tỷ suất lợi nhuận biên của Gói Boost Packs")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Các gói Boost Pack được thiết kế như nguồn thu bổ sung (upsell) với tỷ suất lợi nhuận biên thậm chí còn cao hơn gói Pro cơ bản do không phải phân bổ chi phí duy trì tài khoản:"
    )

    # Table 4: Boost Pack Profits
    table4 = doc.add_table(rows=4, cols=4)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr4 = table4.rows[0].cells
    hdr4[0].text = 'Gói Boost Pack'
    hdr4[1].text = 'Giá bán'
    hdr4[2].text = 'Chi phí API thực tế'
    hdr4[3].text = 'Tỷ suất lợi nhuận (%)'

    for cell in hdr4:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    boost_pricing = [
        ("Boost S (500k tokens / 30 ngày)", "29.000đ", "1.530đ", "94,7%"),
        ("Boost M (1.000k tokens / 60 ngày)", "49.000đ", "3.060đ", "93,7%"),
        ("Boost L (3.000k tokens / 90 ngày)", "119.000đ", "9.180đ", "92,2%")
    ]

    for idx, (p_name, price, cost, margin) in enumerate(boost_pricing):
        row = table4.rows[idx+1].cells
        row[0].text = p_name
        row[1].text = price
        row[2].text = cost
        row[3].text = margin
        
        for c_idx, cell in enumerate(row):
            for p in cell.paragraphs:
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10.5)
        if (idx+1) % 2 == 0:
            for cell in row:
                set_cell_background(cell, "F2F5F8")

    doc.add_page_break()

    # --- Section 3 ---
    h1 = doc.add_paragraph()
    r = h1.add_run("PHẦN III: QUẢN TRỊ RỦI RO CHI PHÍ & TẤM LÁ CHẮN BẢO VỆ HỆ THỐNG")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)

    doc.add_paragraph(
        "Để bảo vệ biên lợi nhuận lý thuyết trên thực tế, hệ thống SmartLife đã thiết lập một hệ thống bảo vệ ba lớp từ máy chủ proxy đến cơ chế kiểm soát dữ liệu đầu vào."
    )

    h2 = doc.add_paragraph()
    r = h2.add_run("1. Chốt chặn Server-Side Quota (Giới hạn Requests & Tokens)")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Hạn mức token hàng tháng (600k tokens) là chốt chặn bảo vệ dài hạn. Tuy nhiên, để ngăn ngừa người dùng chạy script tự động (spam bot) spam hàng vạn requests trong một ngày làm cạn kiệt tài nguyên tức thời, proxy serverless api/gemini.js thực hiện kiểm tra hạn mức hàng ngày:"
    )

    daily_limits = [
        ("Giới hạn lượt gọi (Daily Requests):", " Tối đa 10 yêu cầu/ngày cho mỗi tài khoản Pro. Ngăn chặn triệt để hành vi spam liên tục."),
        ("Giới hạn lưu lượng ngày (Daily Tokens):", " Tối đa 50.000 tokens/ngày. Ngăn chặn hành vi gửi các tài liệu quá lớn liên tục trong ngày làm cạn kiệt quota tháng của chính người dùng.")
    ]

    for title, desc in daily_limits:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(title)
        run_t.font.bold = True
        p.add_run(desc)

    h2 = doc.add_paragraph()
    r = h2.add_run("2. Bảo mật API Keys & Xác thực Token JWT Supabase")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Bảo mật tài sản số của dự án được bảo vệ thông qua kiến trúc Proxy Server:"
    )

    sec_points = [
        ("Ẩn hoàn toàn API keys khỏi Client:", " API key của Gemini không bao giờ được gửi về trình duyệt của người dùng. Mọi truy vấn đều đi qua proxy server api/gemini.js trên Vercel, nơi lưu trữ an toàn các khóa API trong biến môi trường bảo mật."),
        ("Xác thực thông qua Supabase JWT:", " Khách hàng khi gửi request lên proxy bắt buộc phải đính kèm Header Authorization Bearer JWT. Proxy sẽ gọi Supabase auth để xác thực danh tính thực tế, giải mã email và đối soát gói dịch vụ tương ứng, loại bỏ 100% truy cập trái phép từ bên ngoài hệ thống.")
    ]

    for title, desc in sec_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(title)
        run_t.font.bold = True
        p.add_run(desc)

    h2 = doc.add_paragraph()
    r = h2.add_run("3. Biện pháp tối ưu hóa Context giảm thiểu lạm dụng")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Việc tối ưu hóa kỹ thuật mã nguồn (Phase 5) mang lại hiệu quả trực tiếp giúp giảm lượng token tiêu thụ trên mỗi yêu cầu:"
    )

    ctx_points = [
        ("Cắt giảm ngữ cảnh tĩnh không cần thiết:", " Trong tính năng Quick Insight, hệ thống đã loại bỏ hoàn toàn việc nạp thông tin GPA, thời khóa biểu và mục tiêu. Chỉ giữ lại hồ sơ người dùng và dữ liệu thu chi tháng hiện tại, giúp giảm lượng token đầu vào trung bình từ ~5,000 xuống dưới ~1,200 tokens (Tiết kiệm 76% chi phí)."),
        ("Giới hạn số lượng giao dịch đối chiếu:", " Số lượng giao dịch tài chính nạp vào context được giảm từ 30 xuống 15 giao dịch gần nhất, giúp giữ độ dài prompt ổn định và tiết kiệm chi phí tích lũy."),
        ("Giới hạn chiều dài đầu ra (Output Limit):", " Giảm tham số maxOutputTokens từ 2048 xuống 1024 đối với tính năng Quick Insight, bảo đảm AI tập trung trả lời ngắn gọn, súc tích và tiết kiệm token đầu ra đắt đỏ.")
    ]

    for title, desc in ctx_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(title)
        run_t.font.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # --- Section 4 ---
    h1 = doc.add_paragraph()
    r = h1.add_run("PHẦN IV: ĐỊNH VỊ THỊ TRƯỜNG THEO MÔ HÌNH TAM - SAM - SOM")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)

    doc.add_paragraph(
        "Mô hình TAM - SAM - SOM được sử dụng để phân tích quy mô thị trường mục tiêu và định hình tiềm năng tăng trưởng doanh thu thực tế cho dự án SmartLife tại thị trường Việt Nam."
    )

    # Add Diagram Mockup
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_dia = p.add_run(
        "┌──────────────────────────────────────────────────────────┐\n"
        "│           TAM (Total Addressable Market)                 │\n"
        "│       15.000.000 Học sinh, Sinh viên & Người đi làm      │\n"
        "│  └────────────────────────────────────────────────────┘  │\n"
        "│         │    SAM (Serviceable Addressable Market)     │  │\n"
        "│         │  3.500.000 Học sinh THPT & Sinh viên ĐH     │  │\n"
        "│         │  └───────────────────────────────────────┘  │  │\n"
        "│         │       │    SOM (Serviceable Obtainable)  │  │  │\n"
        "│         │       │    70.000 Khách hàng Trả phí     │  │  │\n"
        "│         │       │    (Mục tiêu chiếm lĩnh 2% SAM)  │  │  │\n"
        "└─────────┴───────┴──────────────────────────────────┴──┴──┘"
    )
    r_dia.font.name = 'Consolas'
    r_dia.font.size = Pt(10)
    r_dia.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph() # Spacer

    h2 = doc.add_paragraph()
    r = h2.add_run("1. TAM (Total Addressable Market) - Thị trường tổng thể")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Thị trường tổng thể của SmartLife bao gồm toàn bộ lực lượng học sinh THCS, THPT, sinh viên các trường Đại học/Cao đẳng và người đi làm trẻ tuổi có nhu cầu quản lý hiệu suất cá nhân, theo dõi tài chính và lập lộ trình phát triển bản thân tại Việt Nam. Quy mô ước tính đạt khoảng 15.000.000 người."
    )

    h2 = doc.add_paragraph()
    r = h2.add_run("2. SAM (Serviceable Addressable Market) - Thị trường tiếp cận")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Thị trường có thể tiếp cận hiệu quả của SmartLife tập trung vào học sinh THPT và sinh viên Đại học tại các thành phố lớn (Hà Nội, TP.HCM, Đà Nẵng, Cần Thơ,...). Đây là nhóm đối tượng có độ số hóa cao, sử dụng smartphone/laptop hàng ngày, có thói quen thanh toán qua ví điện tử/ngân hàng trực tuyến và có nhu cầu tự thân cao trong việc tối ưu hóa học tập (GPA) và quản lý tài chính sinh viên. Quy mô ước tính đạt khoảng 3.500.000 người."
    )

    h2 = doc.add_paragraph()
    r = h2.add_run("3. SOM (Serviceable Obtainable Market) - Thị trường mục tiêu thực tế")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "SOM là mục tiêu thực tế của SmartLife trong vòng 2-3 năm đầu hoạt động sau khi triển khai các chiến dịch marketing mục tiêu. SmartLife hướng tới chiếm lĩnh 2% thị phần SAM, tương đương với 70.000 người dùng trả phí (sử dụng gói Pro/Lifetime hoặc mua Boost Packs)."
    )

    h2 = doc.add_paragraph()
    r = h2.add_run("4. Dự phóng Doanh thu & Điểm hòa vốn")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(10)

    doc.add_paragraph(
        "Với giả định ARPU (Doanh thu trung bình trên mỗi người dùng) đạt 200.000 VNĐ / năm (kết hợp các gói 1 tháng, 12 tháng, Lifetime và doanh thu gia tăng từ Boost Packs):\n\n"
        "Doanh thu SOM kỳ vọng = 70.000 người dùng * 200.000 VNĐ = 14.000.000.000 VNĐ (14 tỷ VNĐ / năm).\n\n"
        "Điểm hòa vốn dự án:\n"
        "Chi phí cố định (Hạ tầng, duy trì hệ thống, marketing ban đầu) ước tính khoảng 300 triệu VNĐ/năm. Với biên lợi nhuận gộp dịch vụ đạt trên 94%, dự án chỉ cần tối thiểu 1.600 người dùng Pro hoạt động/năm để đạt điểm hòa vốn tài chính hoàn toàn."
    )

    # Save document
    os.makedirs(os.path.dirname("docs/"), exist_ok=True)
    doc.save("docs/SmartLife_Pricing_Revenue_Strategy.docx")
    print("Document created successfully at docs/SmartLife_Pricing_Revenue_Strategy.docx")

if __name__ == '__main__':
    create_pricing_document()
