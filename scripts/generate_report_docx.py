# -*- coding: utf-8 -*-
# File: scripts/generate_report_docx.py
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_report():
    doc = docx.Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(51, 51, 51)
    style_normal.paragraph_format.line_spacing = 1.3
    style_normal.paragraph_format.space_after = Pt(6)

    PRIMARY_COLOR = RGBColor(31, 78, 121)    # Deep Steel Blue
    SECONDARY_COLOR = RGBColor(79, 93, 115)  # Muted Slate
    HIGHLIGHT_COLOR = RGBColor(184, 15, 10)  # Coral Red

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_run = title_p.add_run("BÁO CÁO TÓM TẮT CHIẾN LƯỢC DOANH THU & ĐỊNH GIÁ SMARTLIFE")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(20)
    sub_run = subtitle_p.add_run("Bản Quy Đổi Từ Báo Cáo Markdown Hệ Thống")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = SECONDARY_COLOR

    # Divider line
    doc.add_paragraph("─" * 58).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Section 1: Pricing Tiers ---
    h1 = doc.add_paragraph()
    r = h1.add_run("1. Tóm tắt Mô hình Định giá Mới (Pricing Tiers)")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)

    doc.add_paragraph(
        "• Gói Free: Không hỗ trợ tính năng AI (Advisor & Quick Insight). Chỉ dùng các tính năng quản lý cơ bản (Lịch trình, todolist, habit, pomodoro, bookmark).\n"
        "• Gói Pro / Lifetime: Hỗ trợ đầy đủ AI với hạn mức tối đa 600.000 tokens/tháng."
    )

    # Pricing Table
    table_p = doc.add_table(rows=5, cols=4)
    table_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table_p.rows[0].cells
    hdr[0].text = "Thời hạn gói Pro"
    hdr[1].text = "Đơn giá thanh toán"
    hdr[2].text = "Đơn giá quy đổi tháng"
    hdr[3].text = "Tỷ lệ tiết kiệm (%)"

    for cell in hdr:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    tiers = [
        ("1 Tháng", "59.000 VNĐ", "59.000 VNĐ / tháng", "Mốc cơ sở (0%)"),
        ("3 Tháng", "139.000 VNĐ", "~46.333 VNĐ / tháng", "Tiết kiệm 21%"),
        ("12 Tháng", "399.000 VNĐ", "~33.250 VNĐ / tháng", "Tiết kiệm 44%"),
        ("Lifetime (Trọn đời)", "899.000 VNĐ", "Sử dụng vĩnh viễn", "Giá trị tối ưu nhất")
    ]

    for idx, (dur, price, conv, save) in enumerate(tiers):
        row = table_p.rows[idx+1].cells
        row[0].text = dur
        row[1].text = price
        row[2].text = conv
        row[3].text = save
        for c_idx, cell in enumerate(row):
            for p in cell.paragraphs:
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
        if (idx+1) % 2 == 0:
            for cell in row:
                set_cell_background(cell, "F2F5F8")

    doc.add_paragraph() # Spacer

    doc.add_paragraph("• Gói AI Boost Packs (Mua thêm khi hết hạn mức tháng):")

    # Boost Pack Table
    table_b = doc.add_table(rows=4, cols=4)
    table_b.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_b = table_b.rows[0].cells
    hdr_b[0].text = "Tên gói Boost"
    hdr_b[1].text = "Tokens bổ sung"
    hdr_b[2].text = "Đơn giá"
    hdr_b[3].text = "Hạn sử dụng"

    for cell in hdr_b:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    boosts = [
        ("AI Boost S", "+500.000 tokens", "29.000 VNĐ", "30 ngày"),
        ("AI Boost M", "+1.000.000 tokens", "49.000 VNĐ", "60 ngày"),
        ("AI Boost L", "+3.000.000 tokens", "119.000 VNĐ", "90 ngày")
    ]

    for idx, (name, tok, price, exp) in enumerate(boosts):
        row = table_b.rows[idx+1].cells
        row[0].text = name
        row[1].text = tok
        row[2].text = price
        row[3].text = exp
        for c_idx, cell in enumerate(row):
            for p in cell.paragraphs:
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
        if (idx+1) % 2 == 0:
            for cell in row:
                set_cell_background(cell, "F2F5F8")

    doc.add_paragraph() # Spacer

    # --- Section 2: Token Cost ---
    h1 = doc.add_paragraph()
    r = h1.add_run("2. Phân tích Chi phí Token AI (Chống Bù Lỗ)")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)

    doc.add_paragraph(
        "Đơn giá Google Gemini 2.5 Flash API:\n"
        "• Input Token: 0.075 USD / 1M tokens (~1.912 VNĐ / triệu tokens)\n"
        "• Output Token: 0.300 USD / 1M tokens (~7.650 VNĐ / triệu tokens)\n"
        "• Đơn giá trung bình (Tỷ lệ 80% input, 20% output): 3.060 VNĐ / 1M tokens."
    )

    h2 = doc.add_paragraph()
    r = h2.add_run("Đối soát Biên lợi nhuận Gộp đối với Gói Pro:")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph(
        "• Chi phí AI tối đa/người dùng/tháng (Sử dụng hết 100% hạn mức 600k tokens):\n"
        "  Chi phí tối đa = 0.6M tokens * 3.060 VNĐ = 1.836 VNĐ/tháng\n"
        "• So với doanh thu gói Pro rẻ nhất (Gói 12 tháng quy đổi là 33.250 VNĐ/tháng):\n"
        "  Biên lợi nhuận gộp tối thiểu = (33.250đ - 1.836đ) / 33.250đ = 94,5%"
    )

    h2 = doc.add_paragraph()
    r = h2.add_run("Đối soát Biên lợi nhuận Gói Boost Packs:")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph(
        "• Gói Boost S (29.000 VNĐ bán / 1.530 VNĐ vốn): Biên lợi nhuận 94,7%\n"
        "• Gói Boost M (49.000 VNĐ bán / 3.060 VNĐ vốn): Biên lợi nhuận 93,7%\n"
        "• Gói Boost L (119.000 VNĐ bán / 9.180 VNĐ vốn): Biên lợi nhuận 92,2%"
    )

    # --- Section 3: Risk Management ---
    h1 = doc.add_paragraph()
    r = h1.add_run("3. Chiến lược Quản trị Rủi ro Hệ thống")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)

    doc.add_paragraph(
        "• Rủi ro spam API: Thiết lập chốt chặn Server-Side Quota giới hạn tối đa 10 requests/ngày và 50.000 tokens/ngày trên mỗi người dùng Pro.\n"
        "• Rủi ro lộ API Keys: Ẩn hoàn toàn khóa API trên Serverless Proxy (api/gemini.js), API key chỉ tồn tại ở môi trường máy chủ bảo mật.\n"
        "• Bảo mật tài khoản: Mọi yêu cầu gọi AI bắt buộc phải xác thực thông qua JWT token do Supabase cung cấp trên Header Authorization.\n"
        "• Tối ưu hóa dung lượng (Context Reduction): Cắt giảm 76% lượng token nạp vào tính năng Quick Insight bằng cách lược bỏ ngữ cảnh tĩnh dư thừa (GPA, thời khóa biểu) và giảm giới hạn lịch sử giao dịch đối chiếu xuống 15 mục."
    )

    # --- Section 4: TAM SAM SOM ---
    h1 = doc.add_paragraph()
    r = h1.add_run("4. Mô hình Phân tích Thị trường TAM - SAM - SOM")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)

    # TAM SAM SOM Table
    table_m = doc.add_table(rows=4, cols=3)
    table_m.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_m = table_m.rows[0].cells
    hdr_m[0].text = "Phân khúc thị trường"
    hdr_m[1].text = "Định nghĩa"
    hdr_m[2].text = "Quy mô ước tính"

    for cell in hdr_m:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    tss = [
        ("TAM (Thị trường tổng thể)", "Toàn bộ học sinh THCS/THPT, sinh viên Đại học & Người đi làm trẻ tại Việt Nam có nhu cầu tự quản lý cuộc sống và học tập.", "15.000.000 người"),
        ("SAM (Thị trường tiếp cận)", "Học sinh THPT và sinh viên tại các thành phố lớn có smartphone/laptop, quen sử dụng ví điện tử/banking chuyển khoản trực tuyến.", "3.500.000 người"),
        ("SOM (Thị trường mục tiêu)", "Mục tiêu thực tế chiếm lĩnh 2% của SAM trong 2-3 năm đầu hoạt động sau khi chạy các chiến dịch Marketing có chủ đích.", "70.000 người dùng trả phí")
    ]

    for idx, (name, dft, size) in enumerate(tss):
        row = table_m.rows[idx+1].cells
        row[0].text = name
        row[1].text = dft
        row[2].text = size
        for c_idx, cell in enumerate(row):
            for p in cell.paragraphs:
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
        if (idx+1) % 2 == 0:
            for cell in row:
                set_cell_background(cell, "F2F5F8")

    doc.add_paragraph() # Spacer

    doc.add_paragraph(
        "• Dự phóng Doanh thu: Với ARPU kỳ vọng đạt 200.000 VNĐ / năm, doanh thu SOM dự phóng đạt 14 tỷ VNĐ/năm.\n"
        "• Điểm hòa vốn: Với chi phí cố định ước tính ~300 triệu VNĐ/năm, dự án chỉ cần đạt tối thiểu 1.600 người dùng Pro hoạt động để bắt đầu sinh lời ròng hoàn toàn."
    )

    # Save
    os.makedirs(os.path.dirname("docs/"), exist_ok=True)
    doc.save("docs/SmartLife_Pricing_Revenue_Strategy_Report.docx")
    print("Report generated successfully at docs/SmartLife_Pricing_Revenue_Strategy_Report.docx")

if __name__ == '__main__':
    generate_report()
